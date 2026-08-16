from __future__ import annotations

import hashlib
import io
import json
import os
import queue
import random
import re
import signal
import socket
import subprocess
import tempfile
import threading
import time
import uuid
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Callable, Optional
from urllib.parse import quote

import numpy as np
import ollama
import requests
import sounddevice as sd
from mlx_lm import load as load_lm, generate as generate_lm
from mlx_lm.sample_utils import make_sampler
from mlx_vlm import load as load_vlm, generate as generate_vlm

from audio_utils import play_audio, select_audio_device, get_device_default_rate
from aec_processor import EchoCanceller

import sys

print(f"[PYTHON] Executable: {sys.executable}")
print(f"[PYTHON] Version: {sys.version}")
print(f"[PYTHON] Path:")
for p in sys.path:
    print(f"    {p}")

try:
    from mflux.models.flux.variants.txt2img.flux import Flux1 as _Flux1Class
    MFLUX_AVAILABLE = True
    print("[MFLUX] Import successful")
except Exception as e:
    _Flux1Class = None
    print(f"[MFLUX] Import failed: {type(e).__name__}: {e}")
    
from schemas import (
    GwenState, InputMode, ModelTierName,
    state_changed_event, chat_message_event, mode_changed_event,
    model_tier_changed_event, file_ready_event, session_started_event,
    session_renamed_event, error_event,
)
import db

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except Exception:
    BS4_AVAILABLE = False

try:
    from ddgs import DDGS
    DDGS_AVAILABLE = True
except Exception:
    DDGS_AVAILABLE = False

try:
    import mlx_whisper
    WHISPER_AVAILABLE = True
except Exception:
    WHISPER_AVAILABLE = False

try:
    from mlx_audio.tts.models.kokoro import KokoroPipeline
    from mlx_audio.tts.utils import load_model as load_tts_model
    TTS_AVAILABLE = True
except Exception:
    TTS_AVAILABLE = False

try:
    import chromadb
    from chromadb.utils import embedding_functions
    CHROMADB_IMPORTABLE = True
except Exception:
    CHROMADB_IMPORTABLE = False

try:
    import fitz
    PYMUPDF_AVAILABLE = True
except Exception:
    PYMUPDF_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

try:
    from PIL import Image as PILImage, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except Exception:
    PIL_AVAILABLE = False

try:
    from mflux.models.flux.variants.txt2img.flux import Flux1 as _Flux1Class
    MFLUX_AVAILABLE = True
    print("[MFLUX] Import successful")
except Exception as e:
    _Flux1Class = None
    MFLUX_AVAILABLE = False
    print(f"[MFLUX] Import failed: {type(e).__name__}: {e}")

try:
    from A14_VoiceGate import VoiceGate
    VOICEGATE_IMPORTABLE = True
except Exception:
    VoiceGate = None
    VOICEGATE_IMPORTABLE = False

try:
    import soundfile as sf
    SOUNDFILE_AVAILABLE = True
except Exception:
    SOUNDFILE_AVAILABLE = False


CHROMA_PATH = os.path.expanduser("~/daughter_ai_chroma")
OUTPUT_DIR = Path("SuperGwen_outputs")
OUTPUT_DIR.mkdir(exist_ok=True)

ENABLE_VOICE_GATE = True
VOICEPRINT_PATH = "Gwen_voiceprint.pt"
VOICE_GATE_THRESHOLD = 0.30
GUARDED_TOOLS = {"quit_app", "update_memory"}

DEFAULT_MODEL_TIER_NAME = "REGULAR"
DEFAULT_INPUT_MODE = InputMode.VOICE

SAMPLE_RATE = 16_000
CHUNK_SIZE = 1_024
SILENCE_THRESHOLD = 0.01
SILENCE_DURATION = 1.5
WAKE_TIMEOUT = 30.0
CACHE_TTL_HOURS = 6
VLM_MAX_TOKENS = 512

INTERRUPT_RMS_THRESHOLD = 0.045
INTERRUPT_SUSTAIN_BLOCKS = 4
INTERRUPT_WARMUP_SECONDS = 0.15

TTS_RATE = 24_000
MIC_BLOCK_SIZE = 512
TTS_SPEED = 1.0

WAKE_WORDS = {"gwen", "hey gwen", "hi gwen", "okay gwen"}
DEACTIVATION_PHRASES = {"deactivate", "go to sleep", "stop listening", "sleep", "go to sleep gwen"}
EXIT_PHRASES = {"exit", "quit", "shut down", "shutdown"}
WHISPER_REPO = "mlx-community/distil-whisper-large-v3"

ACTIVATION_PHRASES = [
    "I'm listening.", "Hey.", "Yeah?", "Uh huh.", "Go ahead.",
    "How can I help?", "What's up?", "Tell me.",
]


def _ts() -> str:
    return datetime.now().strftime("%H:%M:%S")


def log_info(msg: str) -> None:    print(f"[INFO {_ts()}] {msg}")
def log_error(msg: str) -> None:   print(f"[ERROR {_ts()}] {msg}")
def log_success(msg: str) -> None: print(f"[SUCCESS {_ts()}] {msg}")


class ModelTier(Enum):
    VERY_LIGHT = "Very Light"
    LIGHT = "Light"
    REGULAR = "Regular"
    MAX = "MAX"


@dataclass(frozen=True)
class ModelConfig:
    name: str
    backend: str
    family: str
    temperature: float = 0.0
    think: bool = False


MODEL_REGISTRY: dict[ModelTier, ModelConfig] = {
    ModelTier.VERY_LIGHT: ModelConfig(
        name="mlx-community/Qwen2.5-1.5B-Instruct-4bit", backend="mlx_lm", family="qwen", temperature=0.0,
    ),
    ModelTier.LIGHT: ModelConfig(
        name="mlx-community/Qwen3.5-9B-MLX-4bit", backend="mlx_vlm", family="qwen", temperature=0.0,
    ),
    ModelTier.REGULAR: ModelConfig(
        name="gemma4:cloud", backend="ollama", family="gemma", temperature=0.1, think=False,
    ),
    ModelTier.MAX: ModelConfig(
        name="kimi-k2.7-code:cloud", backend="ollama", family="kimi", temperature=0.1, think=False,
    ),
}

REASONING_PATTERNS: dict[str, list[str]] = {
    "qwen": [r"<think>.*?</think>"],
    "gemma": [r"<think>.*?</think>", r"<thinking>.*?</thinking>", r"<reasoning>.*?</reasoning>"],
    "kimi": [r"<think>.*?</think>", r"<thinking>.*?</thinking>", r"<reasoning>.*?</reasoning>"],
}

_WIRE_TO_INTERNAL: dict[ModelTierName, ModelTier] = {
    ModelTierName.VERY_LIGHT: ModelTier.VERY_LIGHT,
    ModelTierName.LIGHT: ModelTier.LIGHT,
    ModelTierName.REGULAR: ModelTier.REGULAR,
    ModelTierName.MAX: ModelTier.MAX,
}
_INTERNAL_TO_WIRE: dict[ModelTier, ModelTierName] = {v: k for k, v in _WIRE_TO_INTERNAL.items()}


TOOL_SCHEMAS: list[dict] = [
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": "Search the web for current, factual, recent, or hard-to-know information. Use this instead of guessing when fresh information is needed.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "The concise web search query."}
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "update_memory",
            "description": "Save a durable fact about the user when the user explicitly asks Gwen to remember it or states a new durable personal fact.",
            "parameters": {
                "type": "object",
                "properties": {
                    "key": {"type": "string", "description": "A specific snake_case memory key."},
                    "value": {"type": "string", "description": "The value to store."}
                },
                "required": ["key", "value"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "generate_image",
            "description": "Generate an image from a natural-language prompt using Gwen's local image-generation model. Use this whenever the user asks to create, generate, draw, make, render, or visualize an image.",
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {"type": "string", "description": "A detailed image-generation prompt preserving the user's requested subject, composition, style, and constraints."},
                    "output_filename": {"type": "string", "description": "Optional output filename ending in .png."},
                    "width": {"type": "integer", "description": "Image width in pixels."},
                    "height": {"type": "integer", "description": "Image height in pixels."},
                    "steps": {"type": "integer", "description": "Number of diffusion inference steps."},
                    "seed": {"type": "integer", "description": "Random seed for reproducible generation."}
                },
                "required": ["prompt"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "understand_image",
            "description": "Inspect and describe or analyze an existing image file using Gwen's vision model.",
            "parameters": {
                "type": "object",
                "properties": {
                    "image_path": {"type": "string", "description": "Path to the image file."},
                    "prompt": {"type": "string", "description": "What to inspect or answer about the image."}
                },
                "required": ["image_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "generate_pdf",
            "description": "Create a PDF document from a title and body supplied by the user.",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "PDF title."},
                    "body": {"type": "string", "description": "PDF body text."},
                    "output_filename": {"type": "string", "description": "Optional output filename ending in .pdf."}
                },
                "required": ["title", "body"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_pdf",
            "description": "Read a PDF from a local path and answer a question or summarize its contents.",
            "parameters": {
                "type": "object",
                "properties": {
                    "pdf_path": {"type": "string", "description": "Path to the PDF file."},
                    "prompt": {"type": "string", "description": "What to extract, summarize, or answer from the PDF."}
                },
                "required": ["pdf_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "generate_docx",
            "description": "Create a Word document from a title and body supplied by the user.",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "Document title."},
                    "body": {"type": "string", "description": "Document body text."},
                    "output_filename": {"type": "string", "description": "Optional output filename ending in .docx."}
                },
                "required": ["title", "body"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_docx",
            "description": "Read a Word document from a local path and answer a question or summarize its contents.",
            "parameters": {
                "type": "object",
                "properties": {
                    "docx_path": {"type": "string", "description": "Path to the Word document."},
                    "prompt": {"type": "string", "description": "What to extract, summarize, or answer from the document."}
                },
                "required": ["docx_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "open_app",
            "description": "Open a macOS application by its application name when the user explicitly asks Gwen to open it.",
            "parameters": {
                "type": "object",
                "properties": {
                    "app": {"type": "string", "description": "macOS application name."}
                },
                "required": ["app"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "quit_app",
            "description": "Quit Gwen when the user explicitly asks Gwen to exit or shut down.",
            "parameters": {
                "type": "object",
                "properties": {},
                "required": []
            }
        }
    },
]

TOOL_NAMES = {tool["function"]["name"] for tool in TOOL_SCHEMAS}


def ollama_available() -> bool:
    try:
        ollama.list()
        return True
    except Exception:
        return False


def network_connection_available(timeout: float = 3.0) -> bool:
    try:
        socket.create_connection(("8.8.8.8", 53), timeout=timeout)
        return True
    except OSError:
        return False


def can_use_ollama_cloud() -> bool:
    return network_connection_available() and ollama_available()


def clean_response(text: str, family: str) -> str:
    for pattern in REASONING_PATTERNS.get(family, []):
        text = re.sub(pattern, "", text, flags=re.DOTALL | re.IGNORECASE)
    return text.strip()


def strip_special_tokens(text: str) -> str:
    text = re.sub(r"<\|im_start\|>\s*assistant\s*", "", text)
    text = re.sub(r"<\|im_end\|>", "", text)
    return text.strip()


def extract_mlx_text(response: Any) -> str:
    return response.text if hasattr(response, "text") else str(response)


def extract_ollama_text(response: Any) -> str:
    if hasattr(response, "message"):
        message = response.message
        if hasattr(message, "content"):
            return message.content or ""
    if isinstance(response, dict):
        message = response.get("message", {})
        if isinstance(message, dict):
            return message.get("content", "") or ""
    return str(response)


def sanitize_for_speech(text: str) -> str:
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    text = re.sub(r"#{1,6}\s*", "", text)
    text = re.sub(r"`{1,3}[^`]*`{1,3}", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"[_~>|*]", " ", text)
    text = text.encode("ascii", errors="ignore").decode("ascii")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def clean_for_display(text: str) -> str:
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)
    text = re.sub(r"<thinking>.*?</thinking>", "", text, flags=re.DOTALL)
    text = re.sub(r"<reasoning>.*?</reasoning>", "", text, flags=re.DOTALL)
    text = strip_special_tokens(text)
    return text.strip()


def resample_audio(audio: np.ndarray, orig_sr: int, target_sr: int) -> np.ndarray:
    if orig_sr == target_sr or audio.size == 0:
        return audio
    ratio = target_sr / orig_sr
    n = int(round(audio.shape[0] * ratio))
    idx = np.arange(n) / ratio
    left = np.floor(idx).astype(np.int64)
    right = np.minimum(left + 1, audio.shape[0] - 1)
    frac = idx - left
    return ((1 - frac) * audio[left] + frac * audio[right]).astype(audio.dtype)


class ModelRouter:
    def __init__(self, mlx_executor: ThreadPoolExecutor):
        self.loaded_models: dict[str, tuple[Any, Any]] = {}
        self.current_tier: Optional[ModelTier] = None
        self._mlx_executor = mlx_executor

    def _load_backend(self, tier: ModelTier) -> bool:
        config = MODEL_REGISTRY[tier]
        if config.backend == "ollama":
            if not can_use_ollama_cloud():
                log_error(f"Cannot reach Ollama cloud for {config.name}.")
                return False
            return True
        if config.name in self.loaded_models:
            return True
        try:
            if config.backend == "mlx_lm":
                model, tokenizer = self._mlx_executor.submit(load_lm, config.name).result()
                self.loaded_models[config.name] = (model, tokenizer)
            elif config.backend == "mlx_vlm":
                model, processor = self._mlx_executor.submit(load_vlm, config.name).result()
                self.loaded_models[config.name] = (model, processor)
            else:
                raise ValueError(f"Unknown backend: {config.backend}")
            log_success(f"Loaded model: {config.name}")
            return True
        except Exception as e:
            log_error(f"Failed to load {config.name}: {e}")
            return False

    def load(self, tier: ModelTier) -> bool:
        ok = self._load_backend(tier)
        if ok:
            self.current_tier = tier
        return ok

    def ensure_loaded(self, tier: ModelTier) -> bool:
        return self._load_backend(tier)

    def unload(self, tier: ModelTier) -> None:
        config = MODEL_REGISTRY[tier]
        if config.name in self.loaded_models:
            del self.loaded_models[config.name]
            try:
                import mlx.core as mx
                mx.metal.clear_cache()
            except Exception:
                pass

    def get_cached(self, tier: ModelTier) -> tuple[Any, Any]:
        config = MODEL_REGISTRY[tier]
        return self.loaded_models[config.name]

    def chat(self, system_prompt: str, history: list[dict], user_input: str, tools: Optional[list[dict]] = None, tool_executor: Optional[Callable[[str, dict], str]] = None, max_tool_rounds: int = 6) -> str:
        if self.current_tier is None:
            raise RuntimeError("No model tier selected. Call load() first.")
        config = MODEL_REGISTRY[self.current_tier]

        if config.backend == "ollama":
            try:
                messages = self._build_ollama_messages(system_prompt, history, user_input)
                return self._generate_ollama_chat(messages, config, tools=tools, tool_executor=tool_executor, max_tool_rounds=max_tool_rounds)
            except RuntimeError as e:
                fallback_cfg = MODEL_REGISTRY[ModelTier.LIGHT]
                log_error(f"{config.name} unavailable ({e}). Falling back to {fallback_cfg.name}.")
                if self.load(ModelTier.LIGHT):
                    prompt = self._build_mlx_prompt(system_prompt, history, user_input)
                    return self._generate_mlx_vlm(prompt, fallback_cfg)
                raise

        if config.backend == "mlx_lm":
            prompt = self._build_mlx_prompt(system_prompt, history, user_input)
            return self._generate_mlx_lm(prompt, config)
        if config.backend == "mlx_vlm":
            prompt = self._build_mlx_prompt(system_prompt, history, user_input)
            return self._generate_mlx_vlm(prompt, config)
        raise ValueError(f"Unknown backend: {config.backend}")

    @staticmethod
    def _build_mlx_prompt(system_prompt: str, history: list[dict], user_input: str) -> str:
        prompt = f"<|im_start|>system\n{system_prompt}<|im_end|>\n"
        for m in history:
            prompt += f"<|im_start|>{m['role']}\n{m['content']}<|im_end|>\n"
        prompt += f"<|im_start|>user\n{user_input}<|im_end|>\n<|im_start|>assistant\n"
        return prompt

    @staticmethod
    def _build_ollama_messages(system_prompt: str, history: list[dict], user_input: str) -> list[dict]:
        messages = [{"role": "system", "content": system_prompt}]
        messages.extend(history)
        messages.append({"role": "user", "content": user_input})
        return messages

    def _generate_mlx_lm(self, prompt: str, config: ModelConfig) -> str:
        model, tokenizer = self.loaded_models[config.name]

        def _run():
            response = generate_lm(model=model, tokenizer=tokenizer, prompt=prompt,
                                    sampler=make_sampler(temp=config.temperature), verbose=False)
            return extract_mlx_text(response)

        raw = self._mlx_executor.submit(_run).result()
        return strip_special_tokens(clean_response(raw, config.family))

    def _generate_mlx_vlm(self, prompt: str, config: ModelConfig) -> str:
        model, processor = self.loaded_models[config.name]

        def _run():
            response = generate_vlm(model=model, processor=processor, prompt=prompt,
                                     temperature=config.temperature, verbose=False)
            return extract_mlx_text(response)

        raw = self._mlx_executor.submit(_run).result()
        return strip_special_tokens(clean_response(raw, config.family))

    def _generate_ollama_chat(self, messages: list[dict], config: ModelConfig, tools: Optional[list[dict]] = None, tool_executor: Optional[Callable[[str, dict], str]] = None, max_tool_rounds: int = 6) -> str:
        try:
            active_tools = tools or []
            for _ in range(max(1, max_tool_rounds)):
                response = ollama.chat(model=config.name, messages=messages, tools=active_tools or None, think=config.think,
                                        options={"temperature": config.temperature})
                message = getattr(response, "message", None)
                tool_calls = getattr(message, "tool_calls", None) if message is not None else None
                if not tool_calls or tool_executor is None:
                    raw = extract_ollama_text(response)
                    return clean_response(raw, config.family)
                messages.append(message)
                for tool_call in tool_calls:
                    function = getattr(tool_call, "function", None)
                    if function is None:
                        continue
                    name = getattr(function, "name", "")
                    arguments = getattr(function, "arguments", {}) or {}
                    if not isinstance(arguments, dict):
                        try:
                            arguments = json.loads(arguments)
                        except Exception:
                            arguments = {}
                    if name not in TOOL_NAMES:
                        result = f"Unknown tool: {name}"
                    else:
                        result = tool_executor(name, arguments)
                    messages.append({"role": "tool", "content": str(result), "tool_name": name})
            response = ollama.chat(model=config.name, messages=messages, tools=active_tools or None, think=config.think,
                                    options={"temperature": config.temperature})
            return clean_response(extract_ollama_text(response), config.family)
        except ollama.ResponseError as e:
            if e.status_code == 401:
                raise RuntimeError("Ollama auth failed. Run: ollama signin") from e
            if e.status_code == 429:
                raise RuntimeError("Ollama cloud limit hit. Try a local tier for now.") from e
            raise RuntimeError(f"Ollama error {e.status_code}: {e}") from e


class GwenEngine:

    def __init__(self):
        self.events: "queue.Queue" = queue.Queue()

        self._mlx_executor = ThreadPoolExecutor(max_workers=1, thread_name_prefix="mlx-worker")

        self.router = ModelRouter(mlx_executor=self._mlx_executor)
        self.echo_canceller = EchoCanceller(stream_delay_ms=50)

        self.state: GwenState = GwenState.IDLE
        self.input_mode: InputMode = DEFAULT_INPUT_MODE
        self.session_id: str = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.chat_history: list[dict] = []

        self._interrupt_event = threading.Event()
        self._stop_event = threading.Event()
        self._pending_text: "queue.Queue[str]" = queue.Queue()

        self.deep_search_mode = False

        self._tts_pipeline = None
        self._small_model = None
        self._small_tokenizer = None
        self.voice_gate = None

        self._init_optional_models()

        db.init_bridge_tables()
        db.create_session(self.session_id)
        self._emit(session_started_event(self.session_id))

        startup_tier = ModelTier[DEFAULT_MODEL_TIER_NAME]
        if not self.router.load(startup_tier):
            log_error(f"Could not load startup tier {startup_tier.value}; falling back to LIGHT")
            self.router.load(ModelTier.LIGHT)

    def _init_optional_models(self) -> None:
        if TTS_AVAILABLE:
            try:
                def _load_tts():
                    voice_model = load_tts_model("prince-canuma/Kokoro-82M")
                    return KokoroPipeline(lang_code="a", model=voice_model, repo_id="prince-canuma/Kokoro-82M")
                self._tts_pipeline = self._mlx_executor.submit(_load_tts).result()
            except Exception as e:
                log_error(f"TTS load failed: {e}")

        try:
            small_repo = MODEL_REGISTRY[ModelTier.VERY_LIGHT].name
            self._small_model, self._small_tokenizer = self._mlx_executor.submit(load_lm, small_repo).result()
            self.router.loaded_models[small_repo] = (self._small_model, self._small_tokenizer)
        except Exception as e:
            log_error(f"Small model unavailable ({e}). Fact extraction disabled.")

        if ENABLE_VOICE_GATE and VOICEGATE_IMPORTABLE:
            vg = VoiceGate(voiceprint_path=VOICEPRINT_PATH, threshold=VOICE_GATE_THRESHOLD)
            if vg.voiceprint is not None:
                self.voice_gate = vg
            else:
                log_error(f"No voiceprint at {VOICEPRINT_PATH}; voice gate disabled, Touch ID fallback expected.")

    def _emit(self, event) -> None:
        self.events.put(event)

    def _set_state(self, new_state: GwenState) -> None:
        if new_state != self.state:
            self.state = new_state
            self._emit(state_changed_event(new_state))

    def set_model_tier(self, tier_name: ModelTierName, source: str = "user_ui") -> str:
        internal_tier = _WIRE_TO_INTERNAL[tier_name]
        if internal_tier == self.router.current_tier:
            return f"Already in {tier_name.value} mode."
        if self.router.load(internal_tier):
            self._emit(model_tier_changed_event(tier_name, source))
            return f"Switched to {tier_name.value} mode."
        return f"Couldn't switch to {tier_name.value} mode — backend unavailable."

    def set_input_mode(self, mode: InputMode, source: str = "user_ui") -> None:
        if mode != self.input_mode:
            self.input_mode = mode
            self._emit(mode_changed_event(mode, source))
            if mode == InputMode.TEXT:
                self._set_state(GwenState.IDLE)

    def submit_text(self, text: str) -> None:
        self._pending_text.put(text)

    def load_session(self, session_id: str) -> dict:
        messages = db.get_session_messages(session_id)
        artifacts = db.get_session_artifacts(session_id)
        self.session_id = session_id
        db.touch_session(session_id)
        self.chat_history = [
            {"role": m["role"], "content": m["display_text"]} for m in messages[-6:]
        ]
        return {"messages": messages, "artifacts": artifacts}

    def start_new_session(self) -> str:
        self.session_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.chat_history = []
        db.create_session(self.session_id)
        self._emit(session_started_event(self.session_id))
        return self.session_id

    def _maybe_generate_session_title(self, user_text: str, assistant_text: str) -> None:
        if self._small_model is None:
            return
        session_id = self.session_id
        prompt = (
            "Summarise the topic of this exchange in 3 to 6 words, like a chat title. "
            "No punctuation at the end, no quotes, no prefix like 'Title:'. "
            "Just the plain title text.\n\n"
            f"User: {user_text[:300]}\n"
            f"Assistant: {assistant_text[:300]}\n\n"
            "Title:"
        )

        def _run():
            response = generate_lm(
                self._small_model, self._small_tokenizer,
                prompt=f"<|im_start|>user\n{prompt}<|im_end|>\n<|im_start|>assistant\n",
                max_tokens=20, sampler=make_sampler(temp=0.2), verbose=False,
            )
            return extract_mlx_text(response)

        try:
            raw = self._mlx_executor.submit(_run).result()
            title = strip_special_tokens(clean_response(raw, "qwen"))
            title = title.strip().strip('"').strip("'").rstrip(".")
            title = re.sub(r"^title:\s*", "", title, flags=re.IGNORECASE).strip()
            if not title:
                return
            if len(title) > 60:
                title = title[:57].rstrip() + "..."
            db.set_session_title(session_id, title)
            self._emit(session_renamed_event(session_id, title))
            log_success(f"[Session] Auto-titled {session_id}: {title!r}")
        except Exception as e:
            log_error(f"[Session] Auto-title failed: {e}")

    def shutdown(self) -> None:
        self._stop_event.set()
        self._interrupt_event.set()

    def _query_hash(self, q: str) -> str:
        return hashlib.md5(q.strip().lower().encode()).hexdigest()

    def get_cached_answer(self, query: str) -> Optional[str]:
        conn = db.get_db()
        row = conn.execute("SELECT answer, cached_at FROM search_cache WHERE query_hash = ?",
                            (self._query_hash(query),)).fetchone()
        conn.close()
        if not row:
            return None
        age_hours = (datetime.now() - datetime.fromisoformat(row["cached_at"])).total_seconds() / 3600
        return None if age_hours > CACHE_TTL_HOURS else row["answer"]

    def cache_answer(self, query: str, answer: str) -> None:
        conn = db.get_db()
        conn.execute("INSERT OR REPLACE INTO search_cache (query_hash, query, answer, cached_at) VALUES (?, ?, ?, ?)",
                     (self._query_hash(query), query, answer, datetime.now().isoformat()))
        conn.commit()
        conn.close()

    def fetch_page_text(self, url: str, max_chars: int = 1500) -> str:
        if not BS4_AVAILABLE:
            return ""
        try:
            r = requests.get(url, timeout=5, headers={"User-Agent": "Mozilla/5.0"})
            soup = BeautifulSoup(r.text, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            return soup.get_text(separator=" ", strip=True)[:max_chars]
        except Exception:
            return ""

    def get_response_with_context(self, user_input: str, context: str) -> str:
        augmented = f"{user_input}\n\n[Web context]:\n{context}"
        system = ("You are a helpful assistant. Answer using ONLY the provided web context. "
                  "Be concise and accurate. If the context doesn't answer the question, say so.")
        raw = self.router.chat(system_prompt=system, history=[], user_input=augmented)
        return clean_for_display(raw)

    def do_web_search(self, query: str, user_input: str, force_fresh: bool = False) -> str:
        if not force_fresh:
            cached = self.get_cached_answer(query)
            if cached:
                return cached
        if not DDGS_AVAILABLE:
            return "Web search isn't available right now — the ddgs package isn't installed."

        variants = [query, f"{query} {datetime.now().strftime('%Y')}",
                    f"{query} today {datetime.now().strftime('%B %d %Y')}"]
        all_context, seen_urls = [], set()
        for q in variants:
            with DDGS() as ddgs:
                try:
                    for r in list(ddgs.text(q, max_results=2)):
                        url = r.get("href", "")
                        if url in seen_urls:
                            continue
                        seen_urls.add(url)
                        snippet = r.get("body", "")
                        page_text = self.fetch_page_text(url) if self.deep_search_mode else ""
                        all_context.append(f"Source: {url}\nSnippet: {snippet}\n{page_text}")
                except Exception:
                    continue
        context = "\n\n---\n\n".join(all_context) if all_context else "No results found."
        answer = self.get_response_with_context(user_input, context)
        self.cache_answer(query, answer)
        return answer

    def _get_vision_model_processor(self):
        if not self.router.ensure_loaded(ModelTier.LIGHT):
            return None, None
        return self.router.get_cached(ModelTier.LIGHT)

    def _vlm_raw(self, prompt: str, images: Optional[list]) -> str:
        model, processor = self._get_vision_model_processor()
        if model is None:
            return "[ERROR] Vision model unavailable."
        try:
            formatted = f"<|im_start|>user\n{prompt}<|im_end|>\n<|im_start|>assistant\n"

            def _run():
                raw = generate_vlm(model, processor, formatted, images, max_tokens=VLM_MAX_TOKENS, verbose=False)
                return extract_mlx_text(raw)

            raw_text = self._mlx_executor.submit(_run).result()
            return strip_special_tokens(clean_response(raw_text, "qwen"))
        except Exception as e:
            return f"[ERROR] Vision inference failed: {e}"

    _imgen_model = None

    def _load_imgen(self) -> bool:
        if GwenEngine._imgen_model is not None:
            return True
        if not MFLUX_AVAILABLE:
            log_error("MFLUX is unavailable. Image generation cannot start.")
            return False
        try:
            GwenEngine._imgen_model = _Flux1Class.from_name(model_name="schnell", quantize=4)
            return True
        except Exception as e:
            log_error(f"Image gen load failed: {e}")
            return False

    def _write_placeholder_image(self, prompt: str, output_path: str) -> bool:
        if not PIL_AVAILABLE:
            return False
        try:
            img = PILImage.new("RGB", (1024, 1024), color=(16, 24, 40))
            draw = ImageDraw.Draw(img)
            font = ImageFont.load_default()
            draw.text((40, 60), "Image generation unavailable", fill=(255, 255, 255), font=font)
            draw.text((40, 160), f"Prompt: {prompt}", fill=(180, 220, 255), font=font)
            img.save(output_path)
            return True
        except Exception as e:
            log_error(f"Placeholder image failed: {e}")
            return False

    def generate_image(self, prompt: str, message_id: str, output_filename: str = "generated.png",
                        width: int = 1024, height: int = 1024, steps: int = 4, seed: int = 42) -> str:
        self.router.unload(ModelTier.LIGHT)
        out_path = str(OUTPUT_DIR / output_filename)
        ok = self._load_imgen()
        if ok:
            try:
                image = GwenEngine._imgen_model.generate_image(
                    seed=seed,
                    prompt=prompt,
                    num_inference_steps=steps,
                    width=width,
                    height=height,
                )
                image.save(path=out_path)
            except Exception as e:
                log_error(f"Image generation failed: {e}")
                ok = self._write_placeholder_image(prompt, out_path)
        else:
            ok = self._write_placeholder_image(prompt, out_path)

        if ok:
            artifact_id = str(uuid.uuid4())
            db.add_artifact(artifact_id, self.session_id, "image", output_filename,
                             message_id=message_id, file_path=out_path)
            self._emit(file_ready_event(self.session_id, out_path, "image", message_id))
        return out_path if ok else ""

    def understand_image(self, image_path: str, prompt: str = "Describe this image in detail.") -> str:
        if not os.path.exists(image_path):
            return f"[ERROR] File not found: {image_path}"
        return self._vlm_raw(prompt, [image_path])

    def read_pdf(self, pdf_path: str, prompt: str = "Summarise the content of this document.") -> str:
        if not PYMUPDF_AVAILABLE:
            return "[ERROR] pymupdf not installed."
        if not os.path.exists(pdf_path):
            return f"[ERROR] File not found: {pdf_path}"
        tmp_paths: list[str] = []
        try:
            doc = fitz.open(pdf_path)
            total_chars = sum(len(p.get_text("text")) for p in doc)
            if total_chars >= 100:
                parts = []
                for i, page in enumerate(doc):
                    text = page.get_text("text").strip()
                    if text:
                        parts.append(f"[Page {i + 1}]\n{text}")
                text = "\n\n".join(parts)
                doc.close()
                return self._vlm_raw(f"{prompt}\n\n---\nDOCUMENT TEXT:\n{text}", None)
            mat = fitz.Matrix(150 / 72, 150 / 72)
            for i, page in enumerate(doc):
                pix = page.get_pixmap(matrix=mat)
                tmp = tempfile.NamedTemporaryFile(suffix=f"_page{i}.png", delete=False, dir=OUTPUT_DIR)
                pix.save(tmp.name)
                tmp.close()
                tmp_paths.append(tmp.name)
            doc.close()
            return self._vlm_raw(prompt, tmp_paths) if tmp_paths else "[ERROR] Could not render PDF pages."
        except Exception as e:
            return f"[ERROR] PDF read failed: {e}"
        finally:
            for p in tmp_paths:
                try: os.unlink(p)
                except OSError: pass

    def generate_pdf(self, title: str, body: str, message_id: str, output_filename: str = "generated.pdf") -> str:
        if not REPORTLAB_AVAILABLE:
            return ""
        out_path = str(OUTPUT_DIR / output_filename)
        try:
            doc = SimpleDocTemplate(out_path, pagesize=A4, leftMargin=2.5 * cm, rightMargin=2.5 * cm,
                                     topMargin=2.5 * cm, bottomMargin=2.5 * cm)
            styles = getSampleStyleSheet()
            story = [Paragraph(title, styles["Title"]), Spacer(1, 0.5 * cm)]
            for para in body.strip().split("\n\n"):
                if para.strip():
                    story.append(Paragraph(para.strip(), styles["BodyText"]))
                    story.append(Spacer(1, 0.3 * cm))
            doc.build(story)
            artifact_id = str(uuid.uuid4())
            db.add_artifact(artifact_id, self.session_id, "pdf", output_filename,
                             message_id=message_id, file_path=out_path)
            self._emit(file_ready_event(self.session_id, out_path, "pdf", message_id))
            return out_path
        except Exception as e:
            log_error(f"PDF generation failed: {e}")
            return ""

    def read_docx(self, docx_path: str, prompt: str = "Summarise the content of this document.") -> str:
        if not DOCX_AVAILABLE:
            return "[ERROR] python-docx not installed."
        if not os.path.exists(docx_path):
            return f"[ERROR] File not found: {docx_path}"
        try:
            doc = DocxDocument(docx_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            tables_text = [" | ".join(c.text.strip() for c in row.cells if c.text.strip())
                            for table in doc.tables for row in table.rows]
            text_body = "\n".join(paragraphs)
            if tables_text:
                text_body += "\n\nTABLES:\n" + "\n".join(t for t in tables_text if t)
            return self._vlm_raw(f"{prompt}\n\nDOCUMENT TEXT:\n{text_body}", None)
        except Exception as e:
            return f"[ERROR] DOCX read failed: {e}"

    def generate_docx(self, title: str, body: str, message_id: str, output_filename: str = "generated.docx") -> str:
        if not DOCX_AVAILABLE:
            return ""
        out_path = str(OUTPUT_DIR / output_filename)
        try:
            doc = DocxDocument()
            doc.add_heading(title, level=0)
            for para in body.strip().split("\n\n"):
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.runs[0].font.size = Pt(11)
            doc.save(out_path)
            artifact_id = str(uuid.uuid4())
            db.add_artifact(artifact_id, self.session_id, "docx", output_filename,
                             message_id=message_id, file_path=out_path)
            self._emit(file_ready_event(self.session_id, out_path, "docx", message_id))
            return out_path
        except Exception as e:
            log_error(f"DOCX generation failed: {e}")
            return ""

    def get_all_facts(self) -> dict:
        conn = db.get_db()
        rows = conn.execute("SELECT key, value FROM user_facts").fetchall()
        conn.close()
        return {r["key"]: r["value"] for r in rows}

    def set_fact(self, key: str, value: str) -> str:
        conn = db.get_db()
        now = datetime.now().isoformat()
        existing = conn.execute("SELECT value FROM user_facts WHERE key = ?", (key,)).fetchone()
        old_value = existing["value"] if existing else None
        if old_value is not None and old_value.strip().lower() == value.strip().lower():
            conn.close()
            log_info(f"[Memory] Skipped no-op update for '{key}' (unchanged).")
            return "unchanged"
        conn.execute("INSERT OR REPLACE INTO user_facts (key, value, updated_at) VALUES (?, ?, ?)", (key, value, now))
        conn.execute("INSERT INTO user_facts_history (key, old_value, new_value, changed_at) VALUES (?, ?, ?, ?)",
                     (key, old_value, value, now))
        conn.commit()
        conn.close()
        log_success(f"[Memory] {'Updated' if old_value else 'Saved'} '{key}': {old_value!r} -> {value!r}")
        return "updated" if old_value else "created"

    def build_global_context(self) -> str:
        facts = self.get_all_facts()
        if not facts:
            return ""
        lines = ["[User context]"] + [f"  {k}: {v}" for k, v in facts.items()]
        return "\n".join(lines)

    def store_semantic_memory(self, text: str, metadata: Optional[dict] = None) -> None:
        pass

    def detect_and_apply_fact_updates(self, user_message: str) -> None:
        if self._small_model is None:
            return
        prompt = (
            "You are a fact extractor. Given a user message, extract any personal facts the user is "
            "stating about themselves. Return a JSON array of {\"key\": ..., \"value\": ...}. "
            "Use snake_case keys. If no facts are stated, return [].\n"
            "Return ONLY the JSON array, nothing else.\n\n"
            f"User message: {user_message}"
        )
        try:
            response = generate_lm(
                self._small_model, self._small_tokenizer,
                prompt=f"<|im_start|>user\n{prompt}<|im_end|>\n<|im_start|>assistant\n",
                max_tokens=200, sampler=make_sampler(temp=0.0), verbose=False,
            )
            raw = clean_response(extract_mlx_text(response), "qwen")
            start, end = raw.find("["), raw.rfind("]")
            if start == -1 or end == -1:
                return
            for f in json.loads(raw[start:end + 1]):
                if "key" in f and "value" in f:
                    self.set_fact(f["key"], f["value"])
        except Exception as e:
            log_error(f"Auto-update failed: {e}")

    def execute_tool(self, name: str, args: dict, user_input: str, message_id: str) -> str:
        name = name.replace(" ", "_").lower()
        args = args if isinstance(args, dict) else {}
        try:
            if name == "web_search":
                query = str(args.get("query", "")).strip()
                if not query:
                    return "A search query was not provided."
                force_fresh = any(w in user_input.lower() for w in ["latest", "again", "refresh", "new search", "today", "current"])
                return self.do_web_search(query, user_input, force_fresh=force_fresh)

            if name == "update_memory":
                key = str(args.get("key", "")).strip()
                value = str(args.get("value", "")).strip()
                if key and value:
                    self.set_fact(key, value)
                    return f"{key.replace('_', ' ')}: {value}"
                return "Memory update requires both a key and a value."

            if name == "generate_image":
                prompt = str(args.get("prompt", "")).strip()
                if not prompt:
                    return "Image generation requires a prompt."
                return self.generate_image(
                    prompt,
                    message_id,
                    str(args.get("output_filename", "generated.png")),
                    int(args.get("width", 1024)),
                    int(args.get("height", 1024)),
                    int(args.get("steps", 4)),
                    int(args.get("seed", 42)),
                )

            if name == "understand_image":
                image_path = str(args.get("image_path", "")).strip()
                if not image_path:
                    return "Image analysis requires an image path."
                return self.understand_image(image_path, str(args.get("prompt", "Describe this image in detail.")))

            if name == "generate_pdf":
                return self.generate_pdf(
                    str(args.get("title", "Untitled")),
                    str(args.get("body", "")),
                    message_id,
                    str(args.get("output_filename", "generated.pdf")),
                )

            if name == "read_pdf":
                return self.read_pdf(
                    str(args.get("pdf_path", "")).strip(),
                    str(args.get("prompt", "Summarise the content of this document.")),
                )

            if name == "generate_docx":
                return self.generate_docx(
                    str(args.get("title", "Untitled")),
                    str(args.get("body", "")),
                    message_id,
                    str(args.get("output_filename", "generated.docx")),
                )

            if name == "read_docx":
                return self.read_docx(
                    str(args.get("docx_path", "")).strip(),
                    str(args.get("prompt", "Summarise the content of this document.")),
                )

            if name == "open_app":
                app = str(args.get("app", "")).strip()
                if not app:
                    return "An application name was not provided."
                subprocess.Popen(["open", "-a", app])
                return f"Opened {app}."

            if name == "quit_app":
                self.shutdown()
                return "Gwen is shutting down."

            return f"Unknown tool: {name}"
        except Exception as e:
            log_error(f"Tool {name} failed: {e}")
            return f"Error executing {name}: {str(e)}"

    @staticmethod
    def _tool_manifest() -> str:
        lines = ["AVAILABLE TOOLS:"]
        for tool in TOOL_SCHEMAS:
            function = tool["function"]
            lines.append(f"- {function['name']}: {function['description']}")
        return "\n".join(lines)

    @staticmethod
    def _looks_like_tool_refusal(text: str) -> bool:
        lowered = text.lower()
        phrases = [
            "i cannot generate",
            "i can't generate",
            "i cannot create",
            "i can't create",
            "i am unable to generate",
            "i'm unable to generate",
            "i am a text-based ai",
            "i'm a text-based ai",
            "i cannot perform that",
            "i can't perform that",
            "i don't have the ability to generate",
            "i do not have the ability to generate",
            "i cannot access",
            "i can't access",
        ]
        return any(phrase in lowered for phrase in phrases)

    @staticmethod
    def _rescue_tool_request(user_input: str) -> Optional[tuple[str, dict]]:
        text = user_input.strip()
        lowered = text.lower()
        image_intent = any(v in lowered for v in ["generate an image", "generate image", "create an image", "create image", "make an image", "make image", "draw an image", "draw image", "render an image", "render image", "generate a picture", "create a picture", "make a picture", "draw a picture"])
        if image_intent:
            return "generate_image", {"prompt": text}
        pdf_intent = any(v in lowered for v in ["generate a pdf", "create a pdf", "make a pdf", "generate pdf", "create pdf", "make pdf"])
        if pdf_intent:
            return "generate_pdf", {"title": "Generated PDF", "body": text}
        docx_intent = any(v in lowered for v in ["generate a word document", "create a word document", "make a word document", "generate a docx", "create a docx", "make a docx"])
        if docx_intent:
            return "generate_docx", {"title": "Generated Document", "body": text}
        app_intent = any(v in lowered for v in ["open the app ", "open app ", "launch the app ", "launch app "])
        if app_intent:
            match = re.search(r"(?:open|launch)(?: the)? app\s+(.+)$", text, flags=re.IGNORECASE)
            if match:
                return "open_app", {"app": match.group(1).strip()}
        return None

    def _build_system_prompt(self) -> str:
        facts_ctx = self.build_global_context()
        tier = self.router.current_tier
        return (
            f"Today's date is {datetime.now().strftime('%B %d, %Y')}. "
            f"You are Gwen, a precise and practical AI assistant. "
            f"You are currently answering in {tier.value} mode, selected manually by the user.\n\n"
            f"{facts_ctx}\n\n"
            "You have real tools available. When a request requires a tool, use the tool instead of refusing, pretending you cannot do it, or merely explaining how it could be done. "
            "Never claim that you cannot generate images, create files, search the web, inspect files, or open applications when the corresponding tool is available. "
            "For image-generation requests, call generate_image. For current or externally verifiable information, call web_search. "
            "For PDF creation, call generate_pdf. For Word document creation, call generate_docx. "
            "For image analysis, call understand_image. For PDF reading, call read_pdf. For Word document reading, call read_docx. "
            "For explicitly requested macOS application launches, call open_app. For explicit shutdown requests, call quit_app.\n\n"
            f"{self._tool_manifest()}\n\n"
            "Use tools whenever appropriate and use the result in your final answer. Do not fabricate tool results. "
            "Do not expose internal tool schemas, tool-routing instructions, hidden reasoning, or chain-of-thought. "
            "Give only the final answer. Do not invent facts. Answer in 1-3 sentences unless asked for more detail or code.\n\n"
            "MEMORY:\n"
            "The [User context] block above, if present, lists everything you already know about the user. Do not ask for or re-save information that is already listed there.\n"
            "When the user states one or more new durable facts about themselves that are NOT already present in [User context], or explicitly asks you to remember something, use update_memory for each distinct fact. "
            "Use a specific snake_case key for each fact. Only store facts about the user themselves, not casual remarks, opinions, or one-off requests. "
            "If a fact already matches the stored context, do not save it again.\n"
            "If the current backend does not expose native tool calling, output the requested tool call as raw JSON using the exact format {\"name\": \"tool_name\", \"arguments\": {...}} and nothing else when a tool is required."
        )

    def handle_turn(self, user_text: str) -> None:
        user_text = user_text.strip()
        if not user_text:
            return

        user_msg_id = str(uuid.uuid4())
        db.append_message(user_msg_id, self.session_id, "user", user_text)
        self._emit(chat_message_event(user_msg_id, self.session_id, "user", user_text))

        if "toggle deep search" in user_text.lower():
            self.deep_search_mode = not self.deep_search_mode
            return

        self._set_state(GwenState.THINKING)
        system_prompt = self._build_system_prompt()
        assistant_msg_id = str(uuid.uuid4())
        attachment: Optional[dict] = None
        file_types = {"generate_image": "image", "generate_pdf": "pdf", "generate_docx": "docx"}
        tool_records: list[dict] = []

        def tool_executor(name: str, args: dict) -> str:
            result = self.execute_tool(name, args, user_text, assistant_msg_id)
            tool_records.append({"name": name, "arguments": args, "result": result})
            return result

        try:
            raw = self.router.chat(
                system_prompt=system_prompt,
                history=self.chat_history[-10:],
                user_input=user_text,
                tools=TOOL_SCHEMAS,
                tool_executor=tool_executor,
                max_tool_rounds=6,
            )

            tool_calls = self._parse_tool_calls(raw)
            if tool_calls:
                for tc in tool_calls:
                    name = tc["name"].replace(" ", "_").lower()
                    tool_executor(name, tc.get("arguments", {}))
                raw = ""

            if not tool_records and self._looks_like_tool_refusal(raw):
                rescue = self._rescue_tool_request(user_text)
                if rescue is not None:
                    tool_executor(rescue[0], rescue[1])
                    raw = ""

            if tool_records:
                for record in tool_records:
                    name = record["name"].replace(" ", "_").lower()
                    result = str(record["result"]).strip()
                    if name in file_types and result and os.path.exists(result):
                        file_type = file_types[name]
                        attachment = {
                            "file_path": result,
                            "file_type": file_type,
                            "filename": os.path.basename(result),
                        }

                if raw.strip():
                    display_text = clean_for_display(raw)
                else:
                    parts = []
                    memory_fragments = []
                    file_messages = []
                    other_results = []
                    for record in tool_records:
                        name = record["name"].replace(" ", "_").lower()
                        result = str(record["result"]).strip()
                        if not result:
                            continue
                        if name == "update_memory":
                            memory_fragments.append(result)
                        elif name in file_types and os.path.exists(result):
                            noun = {"image": "image", "pdf": "PDF", "docx": "Word document"}[file_types[name]]
                            file_messages.append(f"Here's the {noun} you asked for.")
                        elif name == "quit_app":
                            parts.append("Gwen is shutting down.")
                        elif name == "open_app":
                            parts.append(result)
                        elif name in {"web_search", "understand_image", "read_pdf", "read_docx"}:
                            other_results.append(result)
                    if memory_fragments:
                        if len(memory_fragments) == 1:
                            parts.append(f"Got it, I'll remember that {memory_fragments[0]}.")
                        else:
                            parts.append("Got it, I'll remember: " + "; ".join(memory_fragments) + ".")
                    parts.extend(file_messages)
                    parts.extend(other_results)
                    display_text = " ".join(parts).strip() or "Done."
            else:
                display_text = clean_for_display(raw)

            speech_text = sanitize_for_speech(display_text)
        except Exception as e:
            display_text = f"Something went wrong on my end: {e}"
            speech_text = display_text
            self._emit(error_event(str(e)))

        db.append_message(assistant_msg_id, self.session_id, "assistant", display_text, speech_text)
        self._emit(chat_message_event(assistant_msg_id, self.session_id, "assistant", display_text,
                                       speech_text, attachment=attachment))

        self.chat_history.append({"role": "user", "content": user_text})
        self.chat_history.append({"role": "assistant", "content": display_text})

        if len(self.chat_history) == 2:
            self._maybe_generate_session_title(user_text, display_text)

        if self.input_mode == InputMode.VOICE:
            self._set_state(GwenState.SPEAKING)
            try:
                self._speak_interruptible(speech_text)
            except Exception as e:
                log_error(f"speak failed: {e}")

        self._set_state(GwenState.IDLE)

    @staticmethod
    def _parse_tool_calls(text: str) -> list[dict]:
        text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)
        calls: list[dict] = []
        i = 0
        while i < len(text):
            start = text.find("{", i)
            if start == -1:
                break
            depth = 0
            end = None
            for j, ch in enumerate(text[start:], start):
                if ch == "{":
                    depth += 1
                elif ch == "}":
                    depth -= 1
                    if depth == 0:
                        end = j
                        break
            if end is None:
                break
            try:
                data = json.loads(text[start:end + 1])
                if isinstance(data, dict) and "name" in data and "arguments" in data:
                    calls.append(data)
            except json.JSONDecodeError:
                pass
            i = end + 1
        return calls

    def _speak_interruptible(self, text: str) -> bool:
        if self._tts_pipeline is None:
            log_info(f"[TTS unavailable] {text}")
            return False

        self._interrupt_event.clear()
        stop_monitor = threading.Event()
        warmup_until = [time.monotonic() + INTERRUPT_WARMUP_SECONDS]
        monitor = threading.Thread(target=self._mic_rms_monitor, args=(stop_monitor, warmup_until), daemon=True)

        self.echo_canceller.start_reference_stream()
        monitor.start()
        sd.stop()
        time.sleep(0.05)

        safe_device = select_audio_device(sd, kind="output")
        dev_rate = get_device_default_rate(sd, safe_device, fallback=TTS_RATE)

        try:
            for sentence in re.split(r"(?<=[.!?])\s+", text):
                sentence = sentence.strip()
                if not sentence:
                    continue

                def _generate_chunks(_sentence=sentence):
                    return list(self._tts_pipeline(_sentence, voice="af_heart", speed=TTS_SPEED))

                chunks = self._mlx_executor.submit(_generate_chunks).result()

                for _, _, audio in chunks:
                    samples = np.array(audio[0], dtype=np.float32)
                    self.echo_canceller.push_reference(samples, source_sr=TTS_RATE)

                    if dev_rate != TTS_RATE:
                        samples = resample_audio(samples, TTS_RATE, dev_rate)

                    warmup_until[0] = time.monotonic() + INTERRUPT_WARMUP_SECONDS
                    try:
                        play_audio(sd, samples, dev_rate, device_index=safe_device)
                    except Exception as e:
                        log_error(f"play_audio failed: {e}")
                        safe_device = select_audio_device(sd, kind="output")
                        dev_rate = get_device_default_rate(sd, safe_device, fallback=TTS_RATE)
                    if self._interrupt_event.is_set():
                        return True
        finally:
            self.echo_canceller.stop_reference_stream()
            stop_monitor.set()
            monitor.join(timeout=1.0)
            sd.stop()
            interrupted = self._interrupt_event.is_set()
            self._interrupt_event.clear()
        return interrupted

    def _mic_rms_monitor(self, stop_event: threading.Event, warmup_until: list) -> None:
        consecutive_over = 0

        def _callback(indata, frames, time_info, status):
            nonlocal consecutive_over
            if stop_event.is_set():
                raise sd.CallbackStop()
            mic_block = indata[:, 0] if indata.ndim > 1 else indata
            cleaned = self.echo_canceller.process_mic_frame(mic_block)
            if cleaned.size == 0:
                return
            if time.monotonic() < warmup_until[0]:
                consecutive_over = 0
                return
            rms = float(np.sqrt(np.mean(cleaned ** 2)))
            if rms > INTERRUPT_RMS_THRESHOLD:
                consecutive_over += 1
                if consecutive_over >= INTERRUPT_SUSTAIN_BLOCKS:
                    self._interrupt_event.set()
                    raise sd.CallbackStop()
            else:
                consecutive_over = 0

        try:
            with sd.InputStream(samplerate=SAMPLE_RATE, channels=1, dtype="float32",
                                 blocksize=MIC_BLOCK_SIZE, device=0, callback=_callback):
                while not stop_event.is_set() and not self._interrupt_event.is_set():
                    time.sleep(0.01)
        except (sd.CallbackStop, Exception):
            pass
        finally:
            try:
                sd.stop()
            except Exception:
                pass

    def _transcribe_audio(self, audio_data: np.ndarray) -> str:
        if not WHISPER_AVAILABLE or audio_data.size == 0:
            return ""

        def _run():
            result = mlx_whisper.transcribe(audio_data, path_or_hf_repo=WHISPER_REPO, language="en",
                                             task="transcribe", word_timestamps=True, verbose=False)
            return result.get("text", "")

        return self._mlx_executor.submit(_run).result()

    @staticmethod
    def _is_wake_word(text: str) -> bool:
        t = text.strip().lower()
        return any(w in t for w in WAKE_WORDS) or bool(re.search(r"\bgwen\b", t))

    @staticmethod
    def _is_deactivation(text: str) -> bool:
        return any(p in text.strip().lower() for p in DEACTIVATION_PHRASES)

    @staticmethod
    def _is_exit(text: str) -> bool:
        return any(p in text.strip().lower() for p in EXIT_PHRASES)

    def _listen_for_wake_word(self) -> bool:
        with sd.InputStream(samplerate=SAMPLE_RATE, dtype="float32", channels=1,
                             blocksize=CHUNK_SIZE, device=0) as stream:
            while not self._stop_event.is_set() and self.input_mode == InputMode.VOICE:
                if not self._pending_text.empty():
                    return False

                data, _ = stream.read(CHUNK_SIZE)
                if np.sqrt(np.mean(data ** 2)) <= SILENCE_THRESHOLD:
                    continue
                frames = [data.copy()]
                silent_chunks, total = 0, 1
                max_silent = int(0.7 * SAMPLE_RATE / CHUNK_SIZE)
                max_chunks = int(2.5 * SAMPLE_RATE / CHUNK_SIZE)
                while total < max_chunks:
                    chunk, _ = stream.read(CHUNK_SIZE)
                    frames.append(chunk.copy())
                    total += 1
                    if np.sqrt(np.mean(chunk ** 2)) < SILENCE_THRESHOLD:
                        silent_chunks += 1
                        if silent_chunks >= max_silent:
                            break
                    else:
                        silent_chunks = 0
                text = self._transcribe_audio(np.concatenate(frames).flatten())
                if text.strip() and self._is_wake_word(text):
                    return True
        return False

    def _listen_for_prompt(self) -> Optional[str]:
        deadline = time.time() + WAKE_TIMEOUT
        time.sleep(0.25)
        with sd.InputStream(samplerate=SAMPLE_RATE, dtype="float32", channels=1,
                             blocksize=CHUNK_SIZE, device=0) as stream:
            while time.time() < deadline:
                if not self._pending_text.empty():
                    return None

                data, _ = stream.read(CHUNK_SIZE)
                if np.sqrt(np.mean(data ** 2)) <= SILENCE_THRESHOLD:
                    continue
                frames = [data.copy()]
                silent_chunks = 0
                max_silent = int(SILENCE_DURATION * SAMPLE_RATE / CHUNK_SIZE)
                while True:
                    chunk, _ = stream.read(CHUNK_SIZE)
                    frames.append(chunk.copy())
                    chunk_rms = np.sqrt(np.mean(chunk ** 2))
                    peak = np.max(np.abs(chunk))
                    if chunk_rms < SILENCE_THRESHOLD and peak < SILENCE_THRESHOLD * 3:
                        silent_chunks += 1
                        if silent_chunks >= max_silent:
                            break
                    else:
                        silent_chunks = 0
                return self._transcribe_audio(np.concatenate(frames).flatten())
        return None

    def voice_gate_verify(self, seconds: float = 3.0) -> bool:
        if self.voice_gate is None or not SOUNDFILE_AVAILABLE:
            return False
        path = str(OUTPUT_DIR / "verify.wav")
        audio = sd.rec(int(seconds * SAMPLE_RATE), samplerate=SAMPLE_RATE, channels=1, dtype="float32", device=0)
        sd.wait()
        sf.write(path, audio, SAMPLE_RATE)
        score, passed = self.voice_gate.verify_file(path)
        log_info(f"Voice gate score={score:.3f} passed={passed}")
        return passed

    def run_voice_loop(self) -> None:
        while not self._stop_event.is_set():
            try:
                pending = self._pending_text.get(timeout=0.1)
                self.handle_turn(pending)
                continue
            except queue.Empty:
                pass

            if self.input_mode != InputMode.VOICE:
                time.sleep(0.1)
                continue

            self._set_state(GwenState.IDLE)
            heard_wake_word = self._listen_for_wake_word()
            if not heard_wake_word:
                continue

            self._set_state(GwenState.LISTENING)
            transcription = self._listen_for_prompt()

            if transcription is None or not transcription.strip():
                self._set_state(GwenState.IDLE)
                continue

            if self._is_exit(transcription):
                self._set_state(GwenState.IDLE)
                self.shutdown()
                return

            if self._is_deactivation(transcription):
                self._set_state(GwenState.IDLE)
                continue

            self.handle_turn(transcription)

    def start(self) -> threading.Thread:
        t = threading.Thread(target=self.run_voice_loop, daemon=True)
        t.start()
        return t